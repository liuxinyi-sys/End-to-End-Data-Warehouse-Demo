#!/usr/bin/env python3
"""Convert supplementary.md to a formatted Word document."""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shd)

def add_code_block(doc, code_text):
    """Add a code block as a shaded paragraph with monospace font."""
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        # Set East Asian font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        # Gray background via paragraph shading
        pPr = p._element.get_or_add_pPr()
        shd = pPr.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F5F5F5',
            qn('w:val'): 'clear'
        })
        pPr.append(shd)

def add_inline_formatted(doc, text, paragraph):
    """Add text with inline markdown formatting (bold, code, links)."""
    # Pattern: **bold**, `code`, [text](url)
    pattern = r'(\*\*(.+?)\*\*)|(`([^`]+)`)|(\[([^\]]+)\]\([^\)]+\))'
    pos = 0
    for match in re.finditer(pattern, text):
        # Add plain text before this match
        if match.start() > pos:
            plain = text[pos:match.start()]
            run = paragraph.add_run(plain)
            run.font.size = Pt(10.5)
        if match.group(1):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.size = Pt(10.5)
        elif match.group(3):  # Code
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif match.group(5):  # Link
            run = paragraph.add_run(match.group(6))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.size = Pt(10.5)
            run.underline = True
        pos = match.end()
    # Add remaining text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(10.5)

def add_table(doc, lines):
    """Parse markdown table lines and add a Word table."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|'):
            cells = [c.strip() for c in line.split('|')]
            # Remove empty first/last from leading/trailing |
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            rows.append(cells)
    
    if not rows:
        return
    
    # Skip separator row (---|---|---)
    data_rows = [rows[0]]  # header
    for r in rows[1:]:
        if all(re.match(r'^[-:]+$', c) for c in r):
            continue
        data_rows.append(r)
    
    num_cols = len(data_rows[0])
    table = doc.add_table(rows=len(data_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            add_inline_formatted(doc, cell_text, p)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if i == 0:  # Header row
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, 'D9E2F3')
    
    doc.add_paragraph()  # Space after table

def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Table detection
        if line.strip().startswith('|') and i + 1 < len(lines) and '|' in lines[i+1] and '---' in lines[i+1]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue
        
        # Headers
        header_match = re.match(r'^(#{1,4})\s+(.+)', line)
        if header_match:
            level = len(header_match.group(1))
            title = header_match.group(2)
            # Remove markdown formatting from headers
            title = re.sub(r'\*\*(.+?)\*\*', r'\1', title)
            title = re.sub(r'`([^`]+)`', r'\1', title)
            title = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', title)
            heading = doc.add_heading(title, level=min(level, 4))
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            # Add a thin border line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): 'CCCCCC'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Blockquote
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            add_inline_formatted(doc, quote_text, p)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.italic = True
            i += 1
            continue
        
        # List items
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.+)', line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            text = list_match.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75 + indent * 0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if re.match(r'\d+\.', marker):
                prefix = marker + ' '
            else:
                prefix = '• '
            run = p.add_run(prefix)
            run.font.size = Pt(10.5)
            add_inline_formatted(doc, text, p)
            i += 1
            continue
        
        # Empty line
        if line.strip() == '':
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_inline_formatted(doc, line.strip(), p)
        i += 1
    
    # Handle any remaining code block
    if in_code_block and code_buffer:
        add_code_block(doc, '\n'.join(code_buffer))
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base, 'docs', 'supplementary.md')
    docx_path = os.path.join(base, 'docs', 'supplementary.docx')
    convert(md_path, docx_path)
